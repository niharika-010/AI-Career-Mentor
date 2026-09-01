import io
import pytest
import docx
from app.services.document_parser import (
    parse_pdf,
    parse_docx,
    parse_txt,
    normalize_document,
    DocumentParsingError,
)
from app.services.security_scanner import security_scanner_service, FileValidationError


def create_sample_pdf_bytes(text: str = "Jane Doe Resume") -> bytes:
    """Helper to generate a valid minimal PDF file byte stream."""
    return (
        b"%PDF-1.4\n"
        b"1 0 obj <</Type /Catalog /Pages 2 0 R>> endobj\n"
        b"2 0 obj <</Type /Pages /Kids [3 0 R] /Count 1>> endobj\n"
        b"3 0 obj <</Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources <</Font <</F1 4 0 R>>>> /Contents 5 0 R>> endobj\n"
        b"4 0 obj <</Type /Font /Subtype /Type1 /BaseFont /Helvetica>> endobj\n"
        b"5 0 obj <</Length 65>> stream\n"
        b"BT /F1 12 Tf 100 700 Td (" + text.encode("utf-8") + b") Tj ET\n"
        b"endstream\nendobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000056 00000 n \n0000000111 00000 n \n0000000212 00000 n \n0000000281 00000 n \n"
        b"trailer <</Size 6 /Root 1 0 R>>\nstartxref\n396\n%%EOF"
    )


def create_sample_docx_bytes(text: str = "John Smith Resume") -> bytes:
    """Helper to generate a valid DOCX document byte stream."""
    doc = docx.Document()
    doc.add_heading(text, level=1)
    doc.add_paragraph("Summary: Experienced Full Stack Software Engineer.")
    doc.add_paragraph("Skills: Python, FastAPI, React, PostgreSQL, Docker")
    doc.add_paragraph("Experience: 5 years of experience building web applications.")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def test_pdf_parsing():
    pdf_bytes = create_sample_pdf_bytes("Jane Doe Resume Text")
    extracted = parse_pdf(pdf_bytes)
    assert "Jane Doe Resume Text" in extracted


def test_docx_parsing():
    docx_bytes = create_sample_docx_bytes("John Smith Resume Text")
    extracted = parse_docx(docx_bytes)
    assert "John Smith Resume Text" in extracted
    assert "Python" in extracted
    assert "FastAPI" in extracted


def test_txt_parsing():
    txt_bytes = "Plain Text Job Description\nRequirements: 3 years Python experience.".encode("utf-8")
    extracted = parse_txt(txt_bytes)
    assert "Plain Text Job Description" in extracted


def test_document_normalization():
    raw_text = """
    Professional Summary
    Senior Software Engineer with 6 years of experience in cloud architectures.
    
    Technical Skills
    Python, FastAPI, TypeScript, React, Docker, PostgreSQL
    
    Work Experience
    Lead Developer at Tech Corp with 6+ years of experience.
    """
    normalized = normalize_document(raw_text, doc_type="resume")
    assert normalized["doc_type"] == "resume"
    assert "Python" in normalized["extracted_skills"]
    assert "FastAPI" in normalized["extracted_skills"]
    assert normalized["experience_years"] == 6.0
    assert normalized["metadata"]["word_count"] > 10
    assert normalized["metadata"]["preprocessed"] is True


def test_security_scanner_validation():
    # Test valid PDF
    pdf_bytes = create_sample_pdf_bytes()
    sanitized, ext = security_scanner_service.validate_file(pdf_bytes, "test_resume.pdf")
    assert sanitized == "test_resume.pdf"
    assert ext == ".pdf"

    # Test invalid extension
    with pytest.raises(FileValidationError, match="Invalid file extension"):
        security_scanner_service.validate_file(pdf_bytes, "malicious.exe")

    # Test invalid magic bytes
    fake_pdf = b"NOT_A_PDF_HEADER_CONTENT"
    with pytest.raises(FileValidationError, match="Invalid PDF file format"):
        security_scanner_service.validate_file(fake_pdf, "fake.pdf")


def test_security_scanner_malware_heuristics():
    # Test DOS executable header detection
    dos_exe = b"MZ\x90\0\x03\0\0\0\x04\0\0\0\xff\xff"
    is_clean, status, reason = security_scanner_service.scan_for_malware(dos_exe, "sample.pdf")
    assert is_clean is False
    assert status == "FLAGGED"
    assert "Potential executable payload" in reason

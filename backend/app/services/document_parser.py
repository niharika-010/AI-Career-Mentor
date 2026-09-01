import io
import re
from typing import Dict, Any, List
import docx
from pypdf import PdfReader
import pdfplumber


class DocumentParsingError(Exception):
    pass


def parse_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF file bytes using pypdf with pdfplumber fallback."""
    extracted_text = ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages_text.append(text)
        extracted_text = "\n\n".join(pages_text)
    except Exception:
        extracted_text = ""

    # Fallback to pdfplumber if pypdf extracts minimal text (e.g. formatted tables/complex PDF)
    if len(extracted_text.strip()) < 50:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                if pages_text:
                    extracted_text = "\n\n".join(pages_text)
        except Exception as e:
            if not extracted_text:
                raise DocumentParsingError(f"Failed to parse PDF document: {str(e)}")

    if not extracted_text.strip():
        raise DocumentParsingError("PDF document contains no readable text.")

    return extracted_text.strip()


def parse_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX file bytes including paragraphs and table contents."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract content from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_content:
                    table_text.append(" | ".join(row_content))

        full_text = "\n".join(paragraphs)
        if table_text:
            full_text += "\n\n--- Tables ---\n" + "\n".join(table_text)

        if not full_text.strip():
            raise DocumentParsingError("DOCX document contains no readable text.")

        return full_text.strip()
    except Exception as e:
        if isinstance(e, DocumentParsingError):
            raise e
        raise DocumentParsingError(f"Failed to parse DOCX document: {str(e)}")


def parse_txt(file_bytes: bytes) -> str:
    """Extract plain text from TXT file bytes."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="replace")

    if not text.strip():
        raise DocumentParsingError("Text document is empty.")

    return text.strip()


def extract_sections_heuristic(clean_text: str) -> Dict[str, str]:
    """Segment document text into normalized sections based on standard headings."""
    known_headers = {
        "summary": ["summary", "profile", "objective", "about me", "professional summary", "overview"],
        "experience": ["experience", "work experience", "employment history", "work history", "professional experience"],
        "education": ["education", "academic background", "qualification", "qualifications"],
        "skills": ["skills", "technical skills", "core competencies", "technologies", "expertise"],
        "projects": ["projects", "personal projects", "key projects"],
        "certifications": ["certifications", "licenses", "courses", "training"],
        "requirements": ["requirements", "minimum requirements", "what we are looking for", "key qualifications"],
        "responsibilities": ["responsibilities", "duties", "what you will do", "key responsibilities"],
    }

    lines = clean_text.split("\n")
    sections: Dict[str, List[str]] = {"general": []}
    current_section = "general"

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        lower_stripped = stripped.lower().rstrip(":")
        matched_header = None

        # Check if line matches a section header
        if len(stripped) < 40:
            for sec_key, keywords in known_headers.items():
                if lower_stripped in keywords:
                    matched_header = sec_key
                    break

        if matched_header:
            current_section = matched_header
            if current_section not in sections:
                sections[current_section] = []
        else:
            sections[current_section].append(stripped)

    return {sec: "\n".join(content) for sec, content in sections.items() if content}


def extract_skills_heuristic(text: str) -> List[str]:
    """Extract common technical skills from parsed document text."""
    common_skills = [
        "Python", "FastAPI", "Django", "Flask", "JavaScript", "TypeScript", "React", "Next.js",
        "Vue", "Node.js", "Express", "HTML", "CSS", "Tailwind CSS", "SQL", "PostgreSQL",
        "MySQL", "MongoDB", "Redis", "Docker", "Kubernetes", "AWS", "GCP", "Azure", "Git",
        "CI/CD", "REST API", "GraphQL", "PyTorch", "TensorFlow", "Scikit-Learn", "Pandas",
        "NumPy", "C++", "Java", "Go", "Rust", "Linux", "System Design", "Microservices"
    ]
    found = []
    text_upper = text.upper()
    for skill in common_skills:
        # Match word boundary
        pattern = r"\b" + re.escape(skill.upper()) + r"\b"
        if re.search(pattern, text_upper):
            found.append(skill)
    return sorted(list(set(found)))


def extract_yoe_heuristic(text: str) -> float:
    """Extract required or candidate years of experience from document text."""
    patterns = [
        r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)(?:\s*of)?\s*experience",
        r"experience\s*(?:of)?\s*(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                return float(match.group(1))
            except ValueError:
                pass
    return 0.0


def normalize_document(raw_text: str, doc_type: str = "resume") -> Dict[str, Any]:
    """
    Normalizes and structures raw text into a standardized document schema.
    Ensures raw documents are preprocessed before sending to Gemini/LLMs.
    """
    # 1. Clean whitespace and remove control characters
    clean = re.sub(r"[\r\t\f\v]", " ", raw_text)
    clean = re.sub(r" +", " ", clean)
    clean = re.sub(r"\n\s*\n+", "\n\n", clean).strip()

    # 2. Compute metadata
    words = clean.split()
    word_count = len(words)
    char_count = len(clean)
    reading_time_min = round(word_count / 200, 1)

    # 3. Section segmentation
    sections = extract_sections_heuristic(clean)

    # 4. Extract heuristic skills and YOE
    extracted_skills = extract_skills_heuristic(clean)
    yoe = extract_yoe_heuristic(clean)

    # 5. Execute Intelligence Engines
    intelligence_dict = {}
    if doc_type == "resume":
        try:
            from app.services.resume_intelligence import resume_intelligence_engine
            parsed_intel = resume_intelligence_engine.parse_text(clean)
            intelligence_dict = parsed_intel.model_dump()
            
            if parsed_intel.skills:
                intel_skill_names = [s.name for s in parsed_intel.skills]
                extracted_skills = sorted(list(set(extracted_skills + intel_skill_names)))
        except Exception as e:
            intelligence_dict = {"error": f"Resume intelligence extraction error: {str(e)}"}
    elif doc_type == "job_description":
        try:
            from app.services.job_intelligence import job_intelligence_engine
            parsed_job_intel = job_intelligence_engine.parse_text(clean)
            intelligence_dict = parsed_job_intel.model_dump()

            if parsed_job_intel.required_skills:
                job_req_names = [s.name for s in parsed_job_intel.required_skills]
                extracted_skills = sorted(list(set(extracted_skills + job_req_names)))
        except Exception as e:
            intelligence_dict = {"error": f"Job intelligence extraction error: {str(e)}"}

    return {
        "doc_type": doc_type,
        "raw_text": raw_text,
        "clean_text": clean,
        "sections": sections,
        "extracted_skills": extracted_skills,
        "experience_years": yoe,
        "intelligence": intelligence_dict,
        "metadata": {
            "word_count": word_count,
            "character_count": char_count,
            "estimated_reading_time_minutes": reading_time_min,
            "section_count": len(sections),
            "preprocessed": True,
        },
    }
